import streamlit as st
import os
from pathlib import Path
import tempfile
from docx import Document
import PyPDF2
import requests
import json

# Load environment variables
from dotenv import load_dotenv
load_dotenv()

# Configure page
st.set_page_config(page_title="GLR Pipeline (Local)", layout="wide")

def extract_text_from_pdf(pdf_file):
    """Extract text from PDF file"""
    try:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        st.error(f"Error extracting PDF: {str(e)}")
        return ""

def call_llm(prompt, api_key):
    """Call Groq API"""
    url = "https://api.groq.com/openai/v1/chat/completions"
    
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json"
    }
    
    data = {
        "model": "llama-3.3-70b-versatile",
        "messages": [
            {"role": "user", "content": prompt}
        ],
        "temperature": 0.2,
        "max_tokens": 8000
    }
    
    try:
        response = requests.post(url, headers=headers, json=data)
        response.raise_for_status()
        return response.json()["choices"][0]["message"]["content"]
    except Exception as e:
        st.error(f"LLM API Error: {str(e)}")
        return None

def extract_template_content(doc_path):
    """Extract full template as text"""
    doc = Document(doc_path)
    full_text = []
    
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
    
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text)
            if row_text:
                full_text.append(" | ".join(row_text))
    
    return "\n".join(full_text)

def analyze_template_structure(doc_path):
    """Analyze template structure and formatting"""
    doc = Document(doc_path)
    structure = []
    
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            style = para.style.name if para.style else "Normal"
            text = para.text.strip()
            
            if text.endswith(':'):
                para_type = "LABEL"
            elif text.isupper() or any(word in text for word in ["Dwelling", "Roof", "Elevation", "Interior", "Contents", "Review"]):
                para_type = "HEADER"
            elif text.startswith('(') and text.endswith(')'):
                para_type = "INSTRUCTION"
            else:
                para_type = "CONTENT"
            
            structure.append({
                "index": i,
                "type": para_type,
                "style": style,
                "text": text[:100]
            })
    
    return structure

def populate_template_smart(template_path, filled_content):
    """Populate template preserving ALL formatting"""
    doc = Document(template_path)
    
    new_paragraphs = [p.strip() for p in filled_content.split('\n') if p.strip()]
    
    template_paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            template_paragraphs.append(para)
    
    for i, para in enumerate(template_paragraphs):
        if i < len(new_paragraphs):
            for run in para.runs:
                run.text = ""
            
            if para.runs:
                para.runs[0].text = new_paragraphs[i]
            else:
                para.add_run(new_paragraphs[i])
    
    if len(new_paragraphs) > len(template_paragraphs):
        for i in range(len(template_paragraphs), len(new_paragraphs)):
            doc.add_paragraph(new_paragraphs[i])
    
    return doc

# Main UI
st.title("GLR Pipeline - Insurance Template Automation (Local Testing)")
st.markdown("Automate insurance template filling using photo reports and AI")

# Get API key from environment or user input
default_api_key = os.getenv("GROQ_API_KEY", "")

# Sidebar with API key input
with st.sidebar:
    st.header("⚙️ Configuration")
    
    if default_api_key:
        api_key = default_api_key
        st.success("✅ API Key loaded from .env")
        st.caption("Using environment variable")
    else:
        st.warning("⚠️ API Key Required")
        api_key = st.text_input(
            "Groq API Key",
            type="password",
            help="Get your free API key from console.groq.com"
        )
        if api_key:
            st.success("✅ API Key entered")
    
    st.markdown("---")
    st.markdown("### 📖 Instructions")
    st.markdown("1. **Enter your Groq API key** (or set GROQ_API_KEY in .env)")
    st.markdown("2. **Upload ONE template** (.docx)")
    st.markdown("3. **Upload ONE or MORE photo reports** (.pdf)")
    st.markdown("4. **Click Process Documents**")
    st.markdown("5. **Download completed report**")
    
    st.markdown("---")
    st.info("💡 This is the local testing version. For production, use the deployed app at https://dataman003.streamlit.app/")

# File uploads
st.markdown("### 📤 Upload Files")

col1, col2 = st.columns(2)

with col1:
    st.markdown("#### 📄 Template (Required)")
    template_file = st.file_uploader("Choose template file", type=['docx'], key="template")
    if template_file:
        st.success(f"Template: {template_file.name}")

with col2:
    st.markdown("#### 📸 Photo Reports (Required)")
    photo_reports = st.file_uploader("Choose photo report file(s)", type=['pdf'], accept_multiple_files=True, key="reports")
    if photo_reports:
        st.success(f"{len(photo_reports)} report(s) uploaded")

# Process button
st.markdown("---")
if st.button("🚀 Process Documents", type="primary", disabled=not (template_file and photo_reports and api_key), use_container_width=True):
    
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    try:
        # Save template
        status_text.text("Loading template...")
        progress_bar.progress(10)
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_template:
            tmp_template.write(template_file.read())
            template_path = tmp_template.name
        
        # Analyze template
        status_text.text("Analyzing template structure...")
        progress_bar.progress(15)
        template_content = extract_template_content(template_path)
        template_structure = analyze_template_structure(template_path)
        
        structure_summary = "\n".join([
            f"Line {s['index']}: [{s['type']}] {s['text'][:50]}..." 
            for s in template_structure[:30]
        ])
        
        # Extract from photo reports
        status_text.text(f"Extracting from {len(photo_reports)} report(s)...")
        progress_bar.progress(35)
        all_report_text = ""
        for i, pdf_file in enumerate(photo_reports):
            report_text = extract_text_from_pdf(pdf_file)
            all_report_text += f"\n\n=== REPORT {i+1}: {pdf_file.name} ===\n{report_text}"
        
        # AI learns structure
        status_text.text("AI learning template structure...")
        progress_bar.progress(50)
        
        structure_prompt = f"""You are analyzing an insurance GLR template structure.

TEMPLATE STRUCTURE:
{structure_summary}

FULL TEMPLATE CONTENT:
{template_content[:4000]}

Analyze this template and identify:
1. What sections exist
2. What information needs to be filled
3. What instructional text needs to be replaced
4. The order and hierarchy of sections

Respond with a brief analysis."""

        structure_analysis = call_llm(structure_prompt, api_key)
        
        # Generate content
        status_text.text("Generating completed report...")
        progress_bar.progress(65)
        
        prompt = f"""You are an expert insurance claims adjuster. Complete this GLR based on template and reports.

YOUR UNDERSTANDING:
{structure_analysis}

TEMPLATE:
{template_content}

PHOTO REPORTS:
{all_report_text[:12000]}

YOUR TASK:
Rewrite the ENTIRE GLR by:
1. Filling ALL placeholders with actual data
2. DELETING instructional text in parentheses and replacing with actual values
3. DELETING template instructions
4. Writing detailed, professional descriptions
5. Including specific measurements and details
6. Using proper insurance language

CRITICAL:
- MAINTAIN EXACT SAME STRUCTURE as template
- Extract CORRECT date of loss from reports
- Include full addresses
- Be specific about damage
- Delete ALL instructions
- For sections with no data, write "N/A"
- Keep section headers exactly as in template
- One paragraph per line

Return COMPLETE filled report as plain text."""

        llm_response = call_llm(prompt, api_key)
        progress_bar.progress(90)
        
        if llm_response:
            # Create document
            status_text.text("Creating document...")
            filled_doc = populate_template_smart(template_path, llm_response)
            
            # Save
            output_path = tempfile.NamedTemporaryFile(delete=False, suffix='.docx').name
            filled_doc.save(output_path)
            
            st.success("🎉 Document completed!")
            
            # Download
            st.markdown("### 📥 Download Your Completed Report")
            with open(output_path, 'rb') as f:
                st.download_button(
                    label="⬇️ Download Completed GLR Report (DOCX)",
                    data=f.read(),
                    file_name="completed_glr_report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    type="primary",
                    use_container_width=True
                )
            
            # Preview
            with st.expander("📄 Preview Generated Report"):
                st.text(llm_response[:2000] + "..." if len(llm_response) > 2000 else llm_response)
            
            progress_bar.progress(100)
            status_text.text("✅ Complete!")
            
            # Cleanup
            os.unlink(template_path)
            os.unlink(output_path)
        else:
            st.error("❌ Failed to generate report")
    
    except Exception as e:
        st.error(f"❌ Error: {str(e)}")

st.markdown("---")
st.markdown("Built with Streamlit • Powered by Llama 3.3 70B via Groq")
