import streamlit as st
import os
from pathlib import Path
import tempfile
from docx import Document
import PyPDF2
import requests
import json

# Load environment variables
from dotenv import load_dotenv
load_dotenv()

# Configure page
st.set_page_config(page_title="GLR Pipeline", layout="wide")

# Custom CSS for green download button
st.markdown("""
<style>
    div[data-testid="stDownloadButton"] button {
        background-color: #28a745 !important;
        color: white !important;
    }
    div[data-testid="stDownloadButton"] button:hover {
        background-color: #218838 !important;
    }
</style>
""", unsafe_allow_html=True)

def extract_text_from_pdf(pdf_file):
    """Extract text from PDF file"""
    try:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        st.error(f"Error extracting PDF: {str(e)}")
        return ""

def call_llm(prompt, api_key):
    """Call Groq API"""
    url = "https://api.groq.com/openai/v1/chat/completions"
    
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json"
    }
    
    data = {
        "model": "llama-3.3-70b-versatile",
        "messages": [
            {"role": "user", "content": prompt}
        ],
        "temperature": 0.2,
        "max_tokens": 8000
    }
    
    try:
        response = requests.post(url, headers=headers, json=data)
        response.raise_for_status()
        return response.json()["choices"][0]["message"]["content"]
    except Exception as e:
        st.error(f"LLM API Error: {str(e)}")
        return None

def extract_template_content(doc_path):
    """Extract full template as text"""
    doc = Document(doc_path)
    full_text = []
    
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
    
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text)
            if row_text:
                full_text.append(" | ".join(row_text))
    
    return "\n".join(full_text)

def analyze_template_structure(doc_path):
    """Analyze template structure and formatting"""
    doc = Document(doc_path)
    structure = []
    
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            # Get paragraph info
            style = para.style.name if para.style else "Normal"
            text = para.text.strip()
            
            # Identify if it's a header, label, or content
            if text.endswith(':'):
                para_type = "LABEL"
            elif text.isupper() or any(word in text for word in ["Dwelling", "Roof", "Elevation", "Interior", "Contents", "Review"]):
                para_type = "HEADER"
            elif text.startswith('(') and text.endswith(')'):
                para_type = "INSTRUCTION"
            else:
                para_type = "CONTENT"
            
            structure.append({
                "index": i,
                "type": para_type,
                "style": style,
                "text": text[:100]  # First 100 chars for reference
            })
    
    return structure

def populate_template_smart(template_path, filled_content):
    """Populate template preserving ALL formatting"""
    doc = Document(template_path)
    
    # The filled_content is the complete text - we'll do intelligent replacement
    # Split into paragraphs
    new_paragraphs = [p.strip() for p in filled_content.split('\n') if p.strip()]
    
    # Get all paragraphs from template
    template_paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            template_paragraphs.append(para)
    
    # Replace content paragraph by paragraph, preserving formatting
    for i, para in enumerate(template_paragraphs):
        if i < len(new_paragraphs):
            # Clear the text but keep formatting
            for run in para.runs:
                run.text = ""
            
            # Add new text to first run (preserves formatting)
            if para.runs:
                para.runs[0].text = new_paragraphs[i]
            else:
                para.add_run(new_paragraphs[i])
    
    # If there are more new paragraphs than template paragraphs, add them
    if len(new_paragraphs) > len(template_paragraphs):
        for i in range(len(template_paragraphs), len(new_paragraphs)):
            doc.add_paragraph(new_paragraphs[i])
    
    return doc

# Main UI
st.title("GLR Pipeline - Insurance Template Automation")
st.markdown("Automate insurance template filling using photo reports and AI")

# Get API key from environment (Streamlit secrets)
api_key = os.getenv("GROQ_API_KEY", "")

# Try to get from Streamlit secrets (for deployed app)
if not api_key:
    try:
        api_key = st.secrets.get("GROQ_API_KEY", "")
    except:
        pass

# Sidebar - clean and simple, no API key mention
with st.sidebar:
    st.header("📖 Instructions")
    st.markdown("1. **Upload ONE template** (.docx)")
    st.markdown("2. **Upload ONE or MORE photo reports** (.pdf)")
    st.markdown("3. **Click Process Documents**")
    st.markdown("4. **Download completed report**")
    
    st.markdown("---")
    st.info("💡 The AI will automatically extract information from your photo reports and fill the template fields.")

# File uploads
st.markdown("### Upload Files")

col1, col2 = st.columns(2)

with col1:
    st.markdown("#### Template (Required)")
    template_file = st.file_uploader("Choose template file", type=['docx'], key="template")
    if template_file:
        st.success(f"Template: {template_file.name}")

with col2:
    st.markdown("#### Photo Reports (Required)")
    photo_reports = st.file_uploader("Choose photo report file(s)", type=['pdf'], accept_multiple_files=True, key="reports")
    if photo_reports:
        st.success(f"{len(photo_reports)} report(s) uploaded")

# Process button
st.markdown("---")
if st.button("Process Documents", type="primary", disabled=not (template_file and photo_reports and api_key), use_container_width=True):
    
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    try:
        # Save template
        status_text.text("Loading template...")
        progress_bar.progress(10)
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_template:
            tmp_template.write(template_file.read())
            template_path = tmp_template.name
        
        # Step 1: Analyze template structure
        status_text.text("Analyzing template structure...")
        progress_bar.progress(15)
        template_content = extract_template_content(template_path)
        template_structure = analyze_template_structure(template_path)
        
        # Create structure summary for AI
        structure_summary = "\n".join([
            f"Line {s['index']}: [{s['type']}] {s['text'][:50]}..." 
            for s in template_structure[:30]  # First 30 lines
        ])
        
        # Extract from photo reports
        status_text.text(f"Extracting from {len(photo_reports)} report(s)...")
        progress_bar.progress(35)
        all_report_text = ""
        for i, pdf_file in enumerate(photo_reports):
            report_text = extract_text_from_pdf(pdf_file)
            all_report_text += f"\n\n=== REPORT {i+1}: {pdf_file.name} ===\n{report_text}"
        
        # Step 2: Have AI learn the structure first
        status_text.text("AI learning template structure...")
        progress_bar.progress(50)
        
        structure_prompt = f"""You are analyzing an insurance GLR template structure.

TEMPLATE STRUCTURE:
{structure_summary}

FULL TEMPLATE CONTENT:
{template_content[:4000]}

Analyze this template and identify:
1. What sections exist (e.g., Date of Loss, Insurable Interest, Dwelling Description, etc.)
2. What information needs to be filled in each section
3. What instructional text (in parentheses) needs to be replaced
4. The order and hierarchy of sections

Respond with a brief analysis of the template structure."""

        structure_analysis = call_llm(structure_prompt, api_key)
        
        # Step 3: Generate content based on learned structure
        status_text.text("Generating completed report with AI...")
        progress_bar.progress(65)
        
        prompt = f"""You are an expert insurance claims adjuster. You must complete a General Loss Report (GLR) based on a template and photo inspection reports.

YOUR UNDERSTANDING OF THE TEMPLATE STRUCTURE:
{structure_analysis}

ORIGINAL TEMPLATE (with instructions to follow):
{template_content}

PHOTO INSPECTION REPORTS (your source data):
{all_report_text[:12000]}

YOUR TASK:
Rewrite the ENTIRE GLR report by:
1. Filling in ALL placeholders with actual data from the photo reports
2. DELETING all instructional text in parentheses like "(one story, two story, etc.)" and replacing with actual values
3. DELETING template instructions like "(Put N/A if...)" and "(Be sure to describe...)"
4. Writing detailed, professional descriptions of damage based on the photo reports
5. Including specific measurements, counts, and details from the reports
6. Using proper insurance report language and formatting

CRITICAL RULES:
- MAINTAIN THE EXACT SAME STRUCTURE as the template (same sections, same order)
- Each line should correspond to a paragraph in the template
- Extract the CORRECT date of loss from the reports (not the inspection date)
- Include full address with street, city, state, zip
- Be specific about damage: include counts, locations, measurements
- Delete ALL template instructions and parenthetical options
- Write complete sentences in professional insurance language
- For sections with no data (like Supplement, Priors), write "N/A" or appropriate professional response
- Match the style and detail level of professional insurance adjusters
- Keep section headers exactly as they are in the template

FORMATTING:
- Put each paragraph on a new line
- Maintain the same number of sections as the template
- Keep headers like "DwellingRoof", "Front Elevation", etc. exactly as in template

Return the COMPLETE filled report as plain text, one paragraph per line."""

        llm_response = call_llm(prompt, api_key)
        progress_bar.progress(90)
        
        if llm_response:
            # Create filled document
            status_text.text("Creating document...")
            filled_doc = populate_template_smart(template_path, llm_response)
            
            # Save
            output_path = tempfile.NamedTemporaryFile(delete=False, suffix='.docx').name
            filled_doc.save(output_path)
            
            st.success("Document completed!")
            
            # Download with prominent green button
            st.markdown("### 📥 Download Your Completed Report")
            with open(output_path, 'rb') as f:
                st.download_button(
                    label="⬇️ Download Completed GLR Report (DOCX)",
                    data=f.read(),
                    file_name="completed_glr_report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    type="primary",
                    use_container_width=True
                )
            
            # Show preview
            with st.expander("Preview Generated Report"):
                st.text(llm_response[:2000] + "..." if len(llm_response) > 2000 else llm_response)
            
            progress_bar.progress(100)
            status_text.text("Complete!")
            
            # Cleanup
            os.unlink(template_path)
            os.unlink(output_path)
        else:
            st.error("Failed to generate report")
    
    except Exception as e:
        st.error(f"Error: {str(e)}")

st.markdown("---")
st.markdown("Built with Streamlit - Powered by Llama 3.3 70B via Groq")
